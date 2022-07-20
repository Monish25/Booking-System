import tkinter as tk
from tkinter import messagebox
import mysql.connector  
import os
from mysql.connector import Error

def cost1():
    global b
    a = E3.get()
    b = E5.get()
    cost(a,b)

def cost(a,b):
    global a1
    global f1
    dbcrd = 'db.WDC'
    with open(dbcrd) as f:
        data = f.readlines() 
        uname = data[0].rstrip() 
        pword = data[1].rstrip()
        lchst = data[2].rstrip()
      
    SQLC = mysql.connector.connect(host = lchst,user = uname,passwd = pword,database = 'booking')
    cursor = SQLC.cursor()
    exn = 'SELECT F_Fare * %s from flight where F_No =%s'
    cursor.execute(exn,(a,b,))
    z1 = cursor.fetchall()
    cursor.close()
    
    def unpack(a):
        global a1
        
        (a1) = a
    unpack(*z1)
    
    a2 = str(a1)
    f1 = tk.Tk()
    f1.geometry('200x100')
    fare = 'Your fare is INR'+ a2 +'\nfor Travelling by this Plane'
    L1 = tk.Label(f1,text = fare)
    B1 = tk.Button(f1,text = 'Next',command = payment)
    L1.pack()
    B1.pack()

    f1.mainloop()

    
def prnt():
    import docx
    doc = docx.Document()
    doc.add_heading('Flight Bookings', 0)
    
    D1 = E1.get()
    D2 = E2.get()
    D3 = E3.get()
    D4 = E4.get()
    D5 = E5.get()
    D7 = E7.get()
    
    DT1 = 'DETAILS OF PASSENGER:'
    a = 'Passenger Name :',D1
    b = 'Gender :',D2
    c = 'No of Seats:',D3
    d = 'Date of Travel:',D4
    e = 'Phone Number : ',D5
    f = 'Aadhar :',D7
    
    bdr = '---------------------------------------------------------------'
    DT2 = 'Travel Details :'
    a2 = str("".join(map(str, a1)))
    
    
    a3 = str(B2)
    a4 = str(B3)
    a5 = str(B4)
    a6 = str(B7)
    DT3 = 'Fare :',a2,'.INR'
    DT4 = 'Flight Name : ',a3
    DT5 = 'Flight Boarding : ',a4
    DT6 = 'Flight Destination : ',a5
    DT7 = 'Flight Departure Time : ',a6
    
    
    
    
    doc_para = doc.add_paragraph(bdr)
    doc_para = doc.add_paragraph(DT1)
    doc_para = doc.add_paragraph(a)
    doc_para = doc.add_paragraph(b)
    doc_para = doc.add_paragraph(c)
    doc_para = doc.add_paragraph(d)
    doc_para = doc.add_paragraph(e)
    doc_para = doc.add_paragraph(f)
    doc_para = doc.add_paragraph(bdr)
    doc_para = doc.add_paragraph(DT2)
    doc_para = doc.add_paragraph(bdr)
    doc_para = doc.add_paragraph(DT3)
    doc_para = doc.add_paragraph(DT4)
    doc_para = doc.add_paragraph(DT5)
    doc_para = doc.add_paragraph(DT6)
    doc_para = doc.add_paragraph(DT7)
    doc_para = doc.add_paragraph(bdr)
    

    doc.save('Flight.docx')

def data(a):
    global B
    dbcrd = 'db.WDC'
    with open(dbcrd) as f:
        data = f.readlines() 
        uname = data[0].rstrip() 
        pword = data[1].rstrip()
        lchst = data[2].rstrip()
      
    SQLC = mysql.connector.connect(host = lchst,user = uname,passwd = pword,database = 'booking')
    cursor = SQLC.cursor()
    exn = 'SELECT * from flight where F_No =%s'
    cursor.execute(exn,(a,))
    C3 = cursor.fetchall()
    cursor.close()
    
    def unpack(a):
        type(a)
        global B7
        global B2
        global B3
        global B4
        (B1,B2,B3,B4,B5,B6,B7) = a
    unpack(*C3)
    
    
def payment():
    f1.destroy()
    data(b)
    prnt()
    os.system('Payment.py')
    os.system('Flight.Docx')
    kill()

def kill():
    messagebox.showinfo('Complete','The Booking was completed :)')
    D1.destroy()
    
def chk1():
    no = E5.get()
    chk2(no)
def chk2(a):
      
        dbcrd = 'db.WDC'
        with open(dbcrd) as f:
            data = f.readlines() 
            uname = data[0].rstrip() 
            pword = data[1].rstrip()
            lchst = data[2].rstrip()
       
        mySQLConnection = mysql.connector.connect(host = lchst,user = uname,passwd = pword,database = 'booking')
             
        cursor = mySQLConnection.cursor()
        sql_select_query = """select F_No from flight where F_No = %s"""
        a = cursor.execute(sql_select_query,(a,))
        data = cursor.fetchall()
        r =  len(data)
        cursor.close()
        if r == 1:
            cost1()
        else:
           messagebox.showinfo('Error','Not found')
       

h = tk.Tk()
h.geometry('300x100')
h.title('Check')
m = 'The Word Documents Related to the Travel Booking\n Must be Closed before You Can PRoceed,\nIf Closed you can Proceed'
H1 = tk.Label(h,text = m)
H1.pack()
H2 = tk.Button(h,text = 'Ok',command = h.destroy)
H2.pack()
h.mainloop()

global E1
global E2
global E3
global E4
global E6
global E7
global E5
D1=tk.Tk()
D1.title('Flight Booking')
D1.geometry('350x200')

L1 = tk.Label(D1,text='Enter Passenger Name:')
E1 = tk.Entry(D1)
L2 = tk.Label(D1,text='Enter Gender')
E2 = tk.Entry(D1)
L3 = tk.Label(D1,text='Number of seats')
E3 = tk.Entry(D1)
L4 = tk.Label(D1,text='Date:')
E4 = tk.Entry(D1) 
L5 = tk.Label(D1,text='Flight Number')
E5 = tk.Entry(D1)
L6 = tk.Label(D1,text='Phone Number:')
E6 = tk.Entry(D1)
L7 = tk.Label(D1,text='Aadhar')
E7 = tk.Entry(D1)

L1.grid(column = 1,row= 1)
L2.grid(column = 1,row= 2)
L3.grid(column = 1,row= 3)
L4.grid(column = 1,row= 4)
L5.grid(column = 1,row= 5)
L6.grid(column= 1,row= 6)
L7.grid(column= 1,row= 7)

E1.grid(column = 2,row= 1)
E2.grid(column = 2,row= 2)
E3.grid(column = 2,row= 3)
E4.grid(column = 2,row= 4)
E5.grid(column = 2,row= 5)
E6.grid(column = 2,row= 6)
E7.grid(column = 2,row= 7)

B2=tk.Button(D1,text='Next',width=12,height=1,command =chk1)

B2.grid(column=2,row=8 )
D1.mainloop()





